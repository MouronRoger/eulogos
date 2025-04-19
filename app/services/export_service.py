"""Service for exporting texts to various formats."""

import bz2
import gzip
import logging
import os
import re
import tempfile
import uuid
import zipfile
from pathlib import Path
from typing import Any, Dict, List, Optional, Union
from datetime import datetime

import jinja2
from ebooklib import epub
from fastapi import HTTPException
from fpdf import FPDF
from weasyprint import HTML

from app.config import EulogosSettings
from app.models.catalog import Text
from app.services.catalog_service import CatalogService
from app.services.xml_processor_service import XMLProcessorService

# Configure logger
logger = logging.getLogger(__name__)

# Configure Jinja2 templates
templates = jinja2.Environment(loader=jinja2.PackageLoader("app", "templates"), autoescape=jinja2.select_autoescape())


class ExportService:
    """Service for exporting texts to various formats."""

    COMPRESSION_LEVELS = {
        "gzip": {"min": 1, "max": 9, "default": 6},
        "bzip2": {"min": 1, "max": 9, "default": 9},
        "zip": {"min": 0, "max": 9, "default": 6},
    }

    def __init__(
        self,
        catalog_service: CatalogService,
        xml_service: XMLProcessorService,
        settings: Optional[EulogosSettings] = None,
        output_dir: Optional[str] = None,
    ):
        """Initialize the export service.

        Args:
            catalog_service: Service for accessing the catalog
            xml_service: Service for processing XML documents
            settings: Optional application settings
            output_dir: Optional custom output directory
        """
        self.catalog_service = catalog_service
        self.xml_service = xml_service
        self.settings = settings or EulogosSettings()
        self.output_dir = output_dir or "exports"

        # Create output directory
        os.makedirs(self.output_dir, exist_ok=True)

    def export_to_html(self, text_id: str, options: Dict[str, Any] = None) -> Path:
        """Export text to standalone HTML.

        Args:
            text_id: ID of the text to export
            options: Dictionary of export options

        Returns:
            Path to the generated HTML file

        Raises:
            FileNotFoundError: If the text cannot be found
            ValueError: If the ID is invalid
        """
        options = options or {}

        # Get metadata
        metadata = self._get_metadata(text_id)
        
        # Add text_id from options if provided
        if "text_id" in options and options["text_id"]:
            metadata["text_id"] = options["text_id"]

        # Get HTML content using the XML processor
        try:
            # Use the XML service
            html_content = self.xml_service.transform_to_html(text_id, options.get("reference"))
        except Exception as e:
            logger.error(f"Error transforming XML to HTML: {e}")
            raise

        # Add HTML boilerplate
        template = templates.get_template("export_html.jinja2")
        html = template.render(
            title=metadata.get("title", "Untitled"),
            content=html_content,
            metadata=metadata,
            css=options.get("custom_css", ""),
        )

        # Write to file
        output_path = self._get_output_path(text_id, "html", options)
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(html)

        # Compress if requested
        if options.get("compression"):
            output_path = self._compress_file(output_path, options["compression"])

        return Path(output_path)

    def export_to_markdown(self, text_id: str, options: Dict[str, Any] = None) -> Path:
        """Export text to Markdown.

        Args:
            text_id: ID of the text to export
            options: Dictionary of export options

        Returns:
            Path to the generated Markdown file

        Raises:
            FileNotFoundError: If the text cannot be found
            ValueError: If the ID is invalid
            ImportError: If html2text is not installed
        """
        try:
            import html2text
        except ImportError:
            logger.error("html2text is required for Markdown export")
            raise ImportError("html2text is required for Markdown export. Install with: pip install html2text")

        options = options or {}

        # Get HTML content
        html_content = self.xml_service.transform_to_html(text_id, options.get("reference"))

        # Convert HTML to Markdown
        h2t = html2text.HTML2Text()
        h2t.unicode_snob = True  # Preserve unicode characters
        h2t.body_width = 0  # No wrapping
        markdown_content = h2t.handle(html_content)

        # Get metadata
        metadata = self._get_metadata(text_id)

        # Add title and metadata
        if options.get("include_metadata", True):
            markdown_header = f"# {metadata.get('title', 'Untitled')}\n\n"
            if metadata.get("creators"):
                markdown_header += "By " + ", ".join(metadata.get("creators", [])) + "\n\n"
            markdown_content = markdown_header + markdown_content

        # Define output Markdown path
        md_path = self._get_output_path(text_id, "md", options)

        # Write Markdown file
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(markdown_content)

        return Path(md_path)

    def export_to_latex(self, text_id: str, options: Dict[str, Any] = None) -> Path:
        """Export text to LaTeX.

        Args:
            text_id: ID of the text to export
            options: Dictionary of export options

        Returns:
            Path to the generated LaTeX file

        Raises:
            FileNotFoundError: If the text cannot be found
            ValueError: If the ID is invalid
        """
        options = options or {}

        # Get metadata
        metadata = self._get_metadata(text_id)

        # Start building LaTeX document
        latex_content = [
            "\\documentclass{article}",
            "\\usepackage{fontspec}",
            "\\usepackage{polyglossia}",
            "\\setmainlanguage{english}",
            "\\setotherlanguage{greek}",
            "\\setmainfont{Times New Roman}",
            "\\newfontfamily\\greekfont[Script=Greek]{Times New Roman}",
            "\\title{" + metadata.get("title", "Untitled") + "}",
        ]

        # Add authors if available
        if metadata.get("creators"):
            latex_content.append("\\author{" + " \\and ".join(metadata.get("creators", [])) + "}")

        # Begin document
        latex_content.extend(
            [
                "\\begin{document}",
                "\\maketitle",
            ]
        )

        # Load document
        document = self.xml_service.load_document(text_id)

        # If reference specified, limit to that section
        if options.get("reference") and options.get("reference") in document.references:
            pass
        else:
            pass

        # Sort references in document
        sorted_refs = sorted(
            document.references.keys(),
            key=lambda s: [int(t) if t.isdigit() else t.lower() for t in re.split(r"(\d+)", s)],
        )

        # Process sections
        if options.get("reference"):
            # Just process the specified reference
            reference = options.get("reference")
            ref_obj = document.references.get(reference)
            if ref_obj:
                text = ref_obj.text_content
                latex_content.append(f"\\section*{{Section {reference}}}")
                latex_content.append(self._clean_for_latex(text))
                latex_content.append("")
        else:
            # Process all references
            for ref in sorted_refs:
                ref_obj = document.references[ref]

                # Skip if it has a parent (to avoid duplicate content)
                if ref_obj.parent_ref:
                    continue

                # Add section heading
                heading_level = len(ref.split("."))
                if heading_level == 1:
                    latex_content.append(f"\\section*{{Section {ref}}}")
                elif heading_level == 2:
                    latex_content.append(f"\\subsection*{{Section {ref}}}")
                else:
                    latex_content.append(f"\\subsubsection*{{Section {ref}}}")

                # Add text content
                text = ref_obj.text_content
                if text:
                    latex_content.append(self._clean_for_latex(text))
                    latex_content.append("")  # Empty line

        # End document
        latex_content.append("\\end{document}")

        # Define output LaTeX path
        latex_path = self._get_output_path(text_id, "tex", options)

        # Write LaTeX file
        with open(latex_path, "w", encoding="utf-8") as f:
            f.write("\n".join(latex_content))

        return Path(latex_path)

    def export_to_pdf(self, text_id: str, options: Dict[str, Any] = None) -> Path:
        """Export text to PDF.
        
        Args:
            text_id: ID of the text to export
            options: Dictionary of export options
            
        Returns:
            Path to the generated PDF file
        """
        options = options or {}

        # First export to HTML
        html_path = self.export_to_html(text_id, options)

        # Define PDF path
        pdf_path = self._get_output_path(text_id, "pdf", options)

        # Convert to PDF using WeasyPrint
        HTML(html_path).write_pdf(pdf_path)

        # Compress if requested
        if options.get("compression"):
            pdf_path = self._compress_file(pdf_path, options["compression"])

        return Path(pdf_path)

    def export_to_epub(self, text_id: str, options: Dict[str, Any] = None) -> Path:
        """Export text to EPUB.
        
        Args:
            text_id: ID of the text to export
            options: Dictionary of export options
            
        Returns:
            Path to the generated EPUB file
        """
        options = options or {}

        # Get metadata
        metadata = self._get_metadata(text_id)

        # Create EPUB book
        book = epub.EpubBook()

        # Set metadata
        book.set_identifier(str(text_id))
        book.set_title(metadata.get("title", "Untitled"))
        if metadata.get("creators"):
            book.add_author(", ".join(metadata.get("creators", [])))
        book.set_language(metadata.get("language", "en"))

        # Get content
        content = self.xml_service.transform_to_html(text_id, options.get("reference"))

        # Create chapter
        chapter = epub.EpubHtml(title=metadata.get("title", "Untitled"), file_name="text.xhtml", content=content)
        book.add_item(chapter)

        # Add default CSS
        style = epub.EpubItem(
            uid="style", file_name="style.css", media_type="text/css", content=options.get("custom_css", "")
        )
        book.add_item(style)

        # Add navigation
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())

        # Define reading order
        book.spine = ["nav", chapter]

        # Write EPUB file
        epub_path = self._get_output_path(text_id, "epub", options)
        epub.write_epub(str(epub_path), book, {})

        # Compress if requested
        if options.get("compression"):
            epub_path = self._compress_file(epub_path, options["compression"])

        return Path(epub_path)

    def export_to_docx(self, text_id: str, options: Dict[str, Any] = None) -> Path:
        """Export text to DOCX using python-docx.

        Args:
            text_id: ID of the text to export
            options: Dictionary of export options

        Returns:
            Path to the generated DOCX file

        Raises:
            ImportError: If python-docx is not installed
            ValueError: If text cannot be found or processed
        """
        try:
            import docx
            from docx.shared import Pt
        except ImportError:
            raise ImportError("python-docx is required for DOCX export. Install with: pip install python-docx")

        options = options or {}

        # Get metadata
        metadata = self._get_metadata(text_id)

        # Create a new document
        doc = docx.Document()

        # Add title
        doc.add_heading(metadata.get("title", "Untitled"), 0)

        # Add creators if available
        if metadata.get("creators"):
            creators_paragraph = doc.add_paragraph("By " + ", ".join(metadata.get("creators", [])))
            creators_paragraph.style = "Subtitle"

        # Load document
        document = self.xml_service.load_document(text_id)

        # Extract references
        references = self.xml_service.extract_references(document)

        # Sort references naturally
        sorted_refs = sorted(references.keys(), key=lambda x: [int(n) if n.isdigit() else n for n in x.split(".")])

        # Process each reference
        for ref in sorted_refs:
            element = references[ref]
            n_attr = element.get("n")

            # Add section heading
            if n_attr:
                heading_level = len(ref.split("."))
                if heading_level > 9:
                    heading_level = 9  # docx supports heading levels 1-9
                doc.add_heading(f"Section {ref}", heading_level)

            # Add text content
            text = "".join(element.itertext()).strip()
            if text:
                p = doc.add_paragraph(text)
                # Use Unicode font for Greek
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)

        # Define output DOCX path
        docx_path = self._get_output_path(text_id, "docx", options)

        # Save document
        doc.save(docx_path)

        return Path(docx_path)

    def _get_metadata(self, text_id: str) -> Dict[str, Any]:
        """Get metadata for export.

        Args:
            text_id: ID of the text

        Returns:
            Dictionary of metadata
        """
        # Get text metadata from catalog
        text = self.catalog_service.get_text_by_id(text_id)
        if not text:
            logger.warning(f"Text not found for ID: {text_id}")
            return {"title": f"Text {text_id}", "creators": ["Unknown"]}

        # Get author information
        author = None
        if text.author_id:
            author_id = text.author_id
            authors = self.catalog_service.get_all_authors()
            for a in authors:
                if a.id == author_id:
                    author = a
                    break

        # Build metadata
        metadata = {
            "title": getattr(text, "work_name", f"Text {text_id}"),
            "creators": [getattr(author, "name", "Unknown Author")] if author else [],
            "language": getattr(text, "language", "grc"),
            "text_id": text_id
        }

        # Try to load document to get more metadata
        try:
            document = self.xml_service.load_document(text_id)

            # Add document metadata
            if "title" in document.metadata and document.metadata["title"]:
                metadata["title"] = document.metadata["title"]

            if "editor" in document.metadata and document.metadata["editor"]:
                metadata["editor"] = document.metadata["editor"]

            if "translator" in document.metadata and document.metadata["translator"]:
                metadata["translator"] = document.metadata["translator"]

            if "language" in document.metadata and document.metadata["language"]:
                metadata["language"] = document.metadata["language"]

            # Add translator and editor to creators if available
            if "translator" in metadata:
                metadata["creators"].append(f"{metadata['translator']} (trans.)")

            if "editor" in metadata and metadata.get("editor") != metadata.get("translator"):
                metadata["creators"].append(f"{metadata['editor']} (ed.)")
        except Exception as e:
            logger.warning(f"Error loading document metadata: {e}")

        return metadata

    def _get_output_path(self, text_id: str, extension: str, options: Dict[str, Any] = None) -> Path:
        """Get output path for export.

        Args:
            text_id: ID of the text
            extension: File extension (without dot)
            options: Dictionary of export options

        Returns:
            Path to the generated file
        """
        options = options or {}

        # Get text from catalog to use its path for filename generation
        text = self.catalog_service.get_text_by_id(text_id)

        # Use custom filename if provided
        if options.get("filename"):
            filename = options.get("filename")
            if not filename.endswith(f".{extension}"):
                filename = f"{filename}.{extension}"
        elif text and text.path:
            # Generate filename from the text's path
            filename = self._get_export_filename(text, extension)
        else:
            # Fallback to using text_id if no text or path
            filename = f"{text_id}.{extension}"

        return Path(self.output_dir) / filename

    def _get_export_filename(self, text: Text, extension: str) -> str:
        """Get the export filename for a text.

        Args:
            text: The text to export
            extension: The file extension

        Returns:
            The export filename
        """
        if text and text.path:
            return f"{Path(text.path).stem}.{extension}"
        
        # Instead of generating a UUID, use a more deterministic approach
        # This could be the timestamp plus a sanitized version of text_id
        return f"export_text_{datetime.now().strftime('%Y%m%d%H%M%S')}.{extension}"

    def _compress_file(self, input_path: Union[str, Path], compression: str, compression_level: Optional[int] = None) -> str:
        """
        Compress a file using the specified compression format and level.

        Args:
            input_path: Path to the file to compress
            compression: Compression format ('gzip', 'bzip2', or 'zip')
            compression_level: Optional compression level (format-specific ranges)

        Returns:
            str: Path to the compressed file

        Raises:
            HTTPException: If compression fails or invalid parameters are provided
        """
        input_path = str(input_path)
        if not os.path.exists(input_path):
            raise HTTPException(status_code=500, detail=f"Input file not found: {input_path}")

        # Validate compression level
        if compression_level is not None:
            level_range = self.COMPRESSION_LEVELS.get(compression, {})
            if not level_range:
                raise HTTPException(status_code=400, detail=f"Invalid compression format: {compression}")
            if not (level_range["min"] <= compression_level <= level_range["max"]):
                raise HTTPException(
                    status_code=400,
                    detail=f"Compression level for {compression} must be between "
                    f"{level_range['min']} and {level_range['max']}",
                )
        else:
            compression_level = self.COMPRESSION_LEVELS[compression]["default"]

        try:
            if compression == "gzip":
                output_path = f"{input_path}.gz"
                with open(input_path, "rb") as f_in:
                    with gzip.open(output_path, "wb", compresslevel=compression_level) as f_out:
                        f_out.write(f_in.read())

            elif compression == "bzip2":
                output_path = f"{input_path}.bz2"
                with open(input_path, "rb") as f_in:
                    with bz2.open(output_path, "wb", compresslevel=compression_level) as f_out:
                        f_out.write(f_in.read())

            elif compression == "zip":
                output_path = f"{input_path}.zip"
                with zipfile.ZipFile(
                    output_path, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=compression_level
                ) as f_out:
                    f_out.write(input_path, os.path.basename(input_path))

            else:
                raise HTTPException(status_code=400, detail=f"Unsupported compression format: {compression}")

            return output_path

        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Compression failed: {str(e)}")
        finally:
            # Clean up the input file if compression was successful
            if "output_path" in locals() and os.path.exists(output_path):
                try:
                    os.remove(input_path)
                except Exception:
                    pass  # Ignore cleanup errors

    def _get_html_css(self, options: Dict[str, Any] = None) -> str:
        """Get CSS for HTML export.

        Args:
            options: Dictionary of export options

        Returns:
            CSS as string
        """
        options = options or {}

        css = """
        body {
            font-family: "Times New Roman", Times, serif;
            font-size: 16px;
            line-height: 1.6;
            max-width: 800px;
            margin: 0 auto;
            padding: 1em;
        }
        h1, h2, h3, h4, h5, h6 {
            font-family: "Times New Roman", Times, serif;
            margin-top: 1em;
            margin-bottom: 0.5em;
        }
        .reference {
            margin-bottom: 1em;
            position: relative;
        }
        [data-n] {
            font-weight: bold;
            margin-right: 0.5em;
            display: inline-block;
        }
        .line {
            display: block;
            margin-left: 2em;
        }
        """

        # Add custom CSS if provided
        if options.get("custom_css"):
            css += options.get("custom_css")

        return css

    def _clean_for_latex(self, text: str) -> str:
        """Clean text for LaTeX.

        Args:
            text: Text to clean

        Returns:
            Cleaned text
        """
        # Replace LaTeX special characters
        replacements = {
            "&": "\\&",
            "%": "\\%",
            "$": "\\$",
            "#": "\\#",
            "_": "\\_",
            "{": "\\{",
            "}": "\\}",
            "~": "\\textasciitilde{}",
            "^": "\\textasciicircum{}",
            "\\": "\\textbackslash{}",
            "<": "\\textless{}",
            ">": "\\textgreater{}",
        }

        for char, replacement in replacements.items():
            text = text.replace(char, replacement)

        # Handle Greek text
        # Simplified approach: wrap all text in textgreek as it's a Greek text
        if re.search(r"[Α-Ωα-ω]", text):
            return f"\\textgreek{{{text}}}"

        return text

    async def export_text(self, text: str, format: str = "pdf", options: Optional[Dict[str, Any]] = None) -> str:
        """
        Export text to the specified format with optional compression.

        Args:
            text: The text content to export
            format: Output format (pdf, html, etc.)
            options: Dictionary containing export options including compression settings

        Returns:
            str: Path to the exported (and optionally compressed) file

        Raises:
            HTTPException: If export or compression fails
        """
        options = options or {}
        compression = options.get("compression")
        compression_level = options.get("compression_level")

        try:
            # Generate the output file
            output_path = await self._generate_output(text, format, options)

            # Apply compression if specified
            if compression:
                output_path = self._compress_file(str(output_path), compression, compression_level)

            return output_path

        except HTTPException:
            raise  # Re-raise HTTP exceptions
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Export failed: {str(e)}")

    async def _generate_output(self, text: str, format: str, options: Dict[str, Any]) -> str:
        """
        Generate output file in the specified format.

        Args:
            text: The text content to export
            format: Output format (pdf, html, etc.)
            options: Dictionary containing export options

        Returns:
            str: Path to the generated file

        Raises:
            HTTPException: If file generation fails or format is unsupported
        """
        format = format.lower()
        if format not in ["pdf", "html", "txt"]:
            raise HTTPException(status_code=400, detail=f"Unsupported format: {format}")

        # Create temporary output path
        output_dir = Path("temp_exports")
        output_dir.mkdir(exist_ok=True)
        
        # Use timestamp for deterministic filename
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        output_path = output_dir / f"export_{timestamp}.{format}"

        try:
            if format == "pdf":
                await self._generate_pdf(text, output_path, options)
            elif format == "html":
                await self._generate_html(text, output_path, options)
            else:  # txt
                await self._generate_txt(text, output_path, options)

            return str(output_path)

        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to generate {format} file: {str(e)}")

    async def _generate_pdf(self, text: str, output_path: Path, options: Dict[str, Any]) -> None:
        """
        Generate a PDF file from the input text.

        Args:
            text: The text content to convert to PDF
            output_path: Path where the PDF file will be saved
            options: Dictionary containing PDF-specific options like font, size, etc.

        Raises:
            HTTPException: If PDF generation fails
        """
        try:
            # Extract PDF-specific options
            font_size = options.get("font_size", 12)
            font_name = options.get("font_name", "Helvetica")
            margin = options.get("margin", 1.0)  # inches

            # Create PDF document
            doc = FPDF()
            doc.add_page()
            doc.set_font(font_name, size=font_size)

            # Calculate effective page width
            effective_width = doc.w - 2 * (margin * 25.4)  # Convert inches to mm

            # Split text into lines that fit within margins
            lines = self._split_text_to_lines(text, doc, effective_width)

            # Write lines to PDF
            for line in lines:
                doc.cell(0, 10, txt=line, ln=True)

            # Save the PDF
            doc.output(str(output_path))

        except Exception as e:
            raise HTTPException(status_code=500, detail=f"PDF generation failed: {str(e)}")

    def _split_text_to_lines(self, text: str, doc: FPDF, max_width: float) -> List[str]:
        """
        Split text into lines that fit within the specified width.

        Args:
            text: Text to split
            doc: FPDF document instance
            max_width: Maximum width in mm

        Returns:
            List of lines that fit within max_width
        """
        words = text.split()
        lines = []
        current_line = []

        for word in words:
            current_line.append(word)
            line_width = doc.get_string_width(" ".join(current_line))

            if line_width > max_width:
                if len(current_line) > 1:
                    current_line.pop()
                    lines.append(" ".join(current_line))
                    current_line = [word]
                else:
                    # Word is too long, force split it
                    lines.append(word)
                    current_line = []

        if current_line:
            lines.append(" ".join(current_line))

        return lines

    async def _generate_html(self, text: str, output_path: Path, options: Dict[str, Any]) -> None:
        """
        Generate an HTML file from the input text.

        Args:
            text: The text content to convert to HTML
            output_path: Path where the HTML file will be saved
            options: Dictionary containing HTML-specific options like css, title, etc.

        Raises:
            HTTPException: If HTML generation fails
        """
        try:
            # Extract HTML-specific options
            title = options.get("title", "Generated Document")
            css = options.get(
                "css",
                """
                body {
                    font-family: Arial, sans-serif;
                    line-height: 1.6;
                    margin: 40px;
                    max-width: 800px;
                    margin: 40px auto;
                }
                h1 {
                    color: #333;
                    text-align: center;
                }
                p {
                    margin-bottom: 1em;
                    text-align: justify;
                }
            """,
            )

            # Create HTML content with basic styling
            paragraphs = "".join(f"<p>{paragraph}</p>" for paragraph in text.split("\n\n"))
            html_content = f"""
            <!DOCTYPE html>
            <html lang="en">
            <head>
                <meta charset="UTF-8">
                <meta name="viewport" content="width=device-width, initial-scale=1.0">
                <title>{title}</title>
                <style>
                    {css}
                </style>
            </head>
            <body>
                <h1>{title}</h1>
                {paragraphs}
            </body>
            </html>
            """

            # Write HTML file
            output_path.write_text(html_content, encoding="utf-8")

        except Exception as e:
            raise HTTPException(status_code=500, detail=f"HTML generation failed: {str(e)}")

    async def _generate_txt(self, text: str, output_path: Path, options: Dict[str, Any]) -> None:
        """
        Generate a plain text file.

        Args:
            text: The text content to write
            output_path: Path where the text file will be saved
            options: Dictionary containing text-specific options

        Raises:
            HTTPException: If text file generation fails
        """
        try:
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(text)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Text file generation failed: {str(e)}")
